from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = r"C:\Users\theve\Downloads\Nomobug_Capstone_Proposal_CP1_Updated.docx"


BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(20, 35, 50)
MUTED = RGBColor(90, 99, 110)
LIGHT_BLUE = "E8F1FA"
LIGHT_GRAY = "F3F5F7"
PALE_GREEN = "EAF5EC"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_header_repeat(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        hdr = OxmlElement("w:tblHeader")
        hdr.set(qn("w:val"), "true")
        tr_pr.append(hdr)


def add_table(doc, headers, rows, widths=None, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=DARK)
        set_cell_shading(hdr[i], header_fill)
    set_header_repeat(table.rows[0])
    set_row_cant_split(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
        set_row_cant_split(table.rows[-1])
    if widths:
        set_table_width(table, widths)
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for m in ("top", "bottom", "start", "end"):
                node = tc_mar.find(qn(f"w:{m}"))
                if node is None:
                    node = OxmlElement(f"w:{m}")
                    tc_mar.append(node)
                node.set(qn("w:w"), "90")
                node.set(qn("w:type"), "dxa")
    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)
    return p


def add_callout(doc, label, text, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = BLUE
    p.add_run(text)
    doc.add_paragraph()


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK, 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    styles["Title"].font.name = "Calibri Light"
    styles["Title"].font.size = Pt(24)
    styles["Title"].font.color.rgb = BLUE
    styles["Subtitle"].font.name = "Calibri"
    styles["Subtitle"].font.size = Pt(12)
    styles["Subtitle"].font.color.rgb = MUTED


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("CAPSTONE PROJECT PROPOSAL")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = MUTED

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(
        "Weather-Aware Pest Control Decision-Support Dashboard"
    ).bold = True
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Using BigQuery, Spatial Hotspot Detection, Treatment Difficulty Analytics, and Prescriptive Intelligence for Nomobug Pest Control Services"
    )

    doc.add_paragraph()
    meta = [
        ("Project Type", "Bachelor of Computer Science (Data Analytics) Capstone Proposal"),
        ("Company", "Nomobug Pest Control Services"),
        ("Prepared by", "Thevesh A/L Chandran"),
        ("Proposed System", "Web-based analytics and decision-support system"),
        ("Primary Stack", "Google Sheets, Python ETL, BigQuery, Apache Superset, Docker, optional local Ollama"),
        ("Submission Version", "CP1 Proposal Draft, 2026"),
    ]
    add_table(doc, ["Item", "Details"], meta, widths=[1.7, 4.6], header_fill=LIGHT_GRAY)
    add_callout(
        doc,
        "Positioning",
        "This CP1 proposal covers the planning and design of the project. The system is not a CRM; it is a proposed analytics and decision-support system that will help management understand service quality, treatment difficulty, recurrence, area risk, and upsell strategy from existing company data. Full implementation, testing, and presentation will be carried out in CP2.",
        fill=PALE_GREEN,
    )
    doc.add_page_break()


def add_abstract(doc):
    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "This Capstone Project 1 (CP1) proposal presents the planning and design phase for a future Capstone Project 2 (CP2) implementation. Nomobug Pest Control Services currently records valuable operational data in Google Sheets and Google Calendar, including leads, sales, services, warranty claims, refunds, payments, upsell records, technician assignments, and service scheduling. However, the data is distributed across separate files and is mainly used for recording rather than business intelligence. This limits management's ability to identify high-risk service areas, recurring pest problems, weather-associated service risks, treatment difficulty, technician review indicators, and effective upsell opportunities.",
    )
    add_para(
        doc,
        "This project proposes a web-based decision-support dashboard that integrates Nomobug's operational data with weather and location features to support descriptive, predictive, and prescriptive analytics. The system will use Python for extraction, validation, cleaning, transformation, weather enrichment, recurrence analysis, hotspot detection, and scoring; BigQuery as the cloud analytics warehouse; and Apache Superset as the dashboard visualization layer. The proposed analytics methods include day/week/month trend analysis, treatment difficulty scoring, warranty factor categorisation, recurrence window analysis, DBSCAN-based hotspot detection, KDE heatmap visualization, data quality monitoring, scheduled refresh, and lightweight explainable prediction where the dataset size supports it.",
    )
    add_para(
        doc,
        "The project deliberately avoids over-complex machine learning because the available warranty dataset is expected to contain approximately 200+ records. Instead, it prioritizes explainable spatial-temporal analytics, careful weather association analysis, privacy-conscious reporting, recommendation logging, and prescriptive business rules that are suitable for a real SME environment. Future enhancement can add a local Ollama-based AI recommendation service to generate plain-language suggestions without external API cost. The expected CP2 outcome is a company-usable analytics system that helps management improve service planning, reduce repeat-problem blind spots, monitor warranty and refund patterns, and make more data-driven upsell and operational decisions.",
    )


def add_toc(doc):
    add_heading(doc, "Contents", 1)
    for item in [
        "Abstract",
        "Chapter 1: Introduction",
        "Chapter 2: Literature Review",
        "Chapter 3: Methodology",
        "Chapter 4: Work Plan and Timeline",
        "References",
        "Appendix A: Proposed Dashboard Modules",
        "Appendix B: Proposed Data Warehouse Tables",
    ]:
        add_para(doc, item)
    doc.add_page_break()


def add_chapter1(doc):
    add_heading(doc, "Chapter 1: Introduction", 1)
    add_callout(
        doc,
        "CP1 boundary",
        "This document is a CP1 proposal. It defines the problem, literature foundation, proposed architecture, methodology, work plan, expected deliverables, and evaluation approach. The actual data pipeline, BigQuery warehouse, Superset dashboards, analytics scripts, and user testing will be implemented during CP2.",
        fill=LIGHT_BLUE,
    )
    add_heading(doc, "1.1 Background Study", 2)
    add_para(
        doc,
        "Pest control service companies depend heavily on operational data: customer enquiries, confirmed bookings, pest type, premise type, service date, technician or team assignment, payment status, warranty claims, refund records, and follow-up actions. For small and medium-sized enterprises (SMEs), these records are often maintained in spreadsheets because spreadsheets are flexible, familiar, and low-cost. However, as data volume grows, spreadsheet-based workflows become difficult to analyse across business functions.",
    )
    add_para(
        doc,
        "Nomobug Pest Control Services is a pest control SME serving residential and commercial customers across Klang Valley and surrounding areas. The company already has multiple operational data sources, including prospect tracking, sales records, payment received records, payment links, warranty claims, refunds, upsell records, and Google Calendar scheduling. These sources contain enough information to support better management decisions, but they are not currently integrated into a structured analytics platform.",
    )
    add_para(
        doc,
        "Weather and environmental context are also relevant to the pest control domain. Heavy rainfall, humidity, flooding, and temperature variation may be associated with pest movement, customer complaints, or service difficulty. This project will not claim that weather directly causes pest activity or warranty claims. Instead, it will treat weather data as an associated risk indicator that can help management identify areas requiring closer monitoring.",
    )
    add_para(
        doc,
        "The proposed system is therefore a web-based analytics and decision-support dashboard. During CP2, it will process operational data from Google Sheets and Google Calendar, enrich service records with weather and location features, store analytics-ready outputs in BigQuery, and present insights through Apache Superset dashboards. The system is intended to support management decision-making, not replace the company's existing workflow.",
    )

    add_heading(doc, "1.2 Problem Statement", 2)
    add_para(
        doc,
        "Nomobug has useful data, but the data is fragmented and not fully used for decision-making. Management must manually compare separate sheets to understand warranty trends, refund causes, technician review signals, upsell outcomes, and area-based repeat problems. This makes it difficult to answer operational questions quickly and consistently.",
    )
    for item in [
        "Marketing performance is measured mainly by lead source or campaign volume, not by customer quality, payment completion, refund risk, or repeat purchase behaviour.",
        "Warranty claims, refunds, and complimentary service records are not consistently linked with original service data, pest type, technician, area, and weather context.",
        "Treatment difficulty is not measured systematically, so areas with high repeat-problem rate, refund rate, or extra-session demand are not visible in one dashboard.",
        "Pest recurrence timing is not analysed, making it difficult to know when follow-up should happen after treatment.",
        "Weather and flood indicators are not connected to service dates, even though rainfall and humidity may be useful risk indicators for planning.",
        "Upsell performance is difficult to evaluate because some records show package uptake but not rejected or no-response offers.",
        "Management lacks a single browser-based dashboard that summarises service quality, scheduling, financial, warranty, and upsell patterns by day, week, and month.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "1.3 Project Aim", 2)
    add_para(
        doc,
        "The aim of this CP1 proposal is to design a weather-aware pest control decision-support dashboard that will be developed in CP2. The planned system will transform Nomobug's Google Sheets, Google Calendar, service, warranty, refund, payment, and upsell records into actionable business intelligence for service quality, treatment difficulty, recurrence monitoring, area planning, and upsell strategy.",
    )

    add_heading(doc, "1.4 Project Objectives", 2)
    objectives = [
        "To plan the collection and cleaning of Nomobug's operational datasets from Google Sheets, CSV/Excel files, and Google Calendar exports or API access.",
        "To design an analytics-ready BigQuery data warehouse that separates raw staging data, cleaned dimensions/facts, and dashboard-ready marts.",
        "To specify a data quality report that will highlight missing values, duplicate rows, invalid dates, unmatched records, area/postcode match confidence, weather coverage, and calendar matching confidence.",
        "To design weather enrichment for service records using area/postcode coordinates and historical weather features such as rainfall, humidity, and temperature windows.",
        "To design descriptive analytics dashboards for leads, sales, services, scheduling, payments, warranty claims, refunds, and upsell records by day, week, and month.",
        "To define a treatment difficulty method by area, pest type, and team using repeat problem rate, warranty claim rate, refund rate, complimentary service rate, and extra-session indicators.",
        "To define recurrence window analysis by calculating the number of days between original service and repeat/warranty/complimentary signals.",
        "To propose spatial-temporal repeat-problem hotspot detection using DBSCAN-style clustering and KDE heatmap visualisation.",
        "To propose weather-associated warranty and repeat-problem analysis using time-lag features and explainable statistical methods.",
        "To design prescriptive recommendations for service planning and upsell package-fit using transparent scoring rules and management-reviewed assumptions.",
        "To design a recommendation log that records suggested actions, reasons, confidence labels, management response, and later outcome for future improvement.",
        "To plan an evaluation approach for dashboard usability and insight usefulness with Nomobug management or operations users.",
    ]
    for item in objectives:
        add_number(doc, item)

    add_heading(doc, "1.5 Scope and Limitations", 2)
    add_table(
        doc,
        ["In Scope", "Out of Scope / Limitation"],
        [
            ("Data integration from company sheets, calendar, and uploaded files", "Not a full CRM, job dispatch system, WhatsApp automation, or payment gateway"),
            ("Descriptive analytics by day, week, month, area, pest, package, team, and technician", "Results depend on the quality and consistency of existing company records"),
            ("Weather-associated risk indicators using historical rainfall and humidity features", "Weather will be treated as association/risk indicator, not proof of causation"),
            ("DBSCAN/KDE hotspot analysis using coordinates or postcode centroids", "Exact customer addresses are not required in the academic report"),
            ("Treatment difficulty, recurrence window, warranty factor, refund root cause, and upsell package-fit analysis", "Upsell prediction is limited until rejected/no-response offers are recorded"),
            ("Data quality reporting, scheduled refresh, anonymised reporting views, and recommendation logging", "Near-real-time refresh is not streaming real-time and depends on Google Sheets/API access"),
            ("Optional future local Ollama AI explanation layer", "The AI layer will not make final decisions; it will explain analytics-backed recommendations"),
        ],
        widths=[3.1, 3.1],
    )

    add_heading(doc, "1.6 Research Significance", 2)
    add_para(
        doc,
        "The project is significant because it addresses a real computing and business problem faced by an SME: turning fragmented operational records into reliable decision support. It contributes to the company by reducing manual cross-checking, improving visibility into service quality, and supporting better area planning. It contributes academically by combining data warehousing, business intelligence, spatial analysis, weather-associated risk modelling, and prescriptive analytics in a practical SME setting.",
    )
    add_para(
        doc,
        "The sustainability value is also relevant. Better area planning and recurrence monitoring can reduce unnecessary repeat visits, wasted travel, and unplanned service costs. A local or low-cost analytics stack also supports SME adoption by reducing dependence on expensive enterprise software. Privacy is considered through anonymised academic and reporting views, customer IDs, area/postcode-level location, controlled access, and careful wording for technician-related review indicators.",
    )


def add_chapter2(doc):
    add_heading(doc, "Chapter 2: Literature Review", 1)
    add_para(
        doc,
        "This literature review focuses on recent work published from 2021 to 2026. The review is organised around five themes that directly support the proposed system: business intelligence for SMEs, pest control and weather-associated risk, spatial-temporal hotspot detection, lightweight prediction and explainability, and AI-assisted decision support.",
    )

    add_heading(doc, "2.1 Business Intelligence and SME Decision Support", 2)
    add_para(
        doc,
        "Ragazou et al. (2023) argue that SMEs can benefit from business intelligence because BI converts scattered operational data into information that supports faster and more competitive decision-making. Their work is relevant to Nomobug because the company already collects operational data but lacks a structured analytics model. Gurcan et al. (2023) further show that modern BI research increasingly emphasises data mining, big data analytics, visualisation, and AI applications, which supports the proposed combination of descriptive, predictive, and prescriptive dashboard modules.",
    )
    add_para(
        doc,
        "For this project, the literature supports treating BI as more than simple reporting. The dashboard should not only show totals, but also connect business questions across sources: which ad sources produce better customers, which areas generate repeat problems, which service types cause higher refunds, and which package-fit recommendations are supported by historical patterns.",
    )

    add_heading(doc, "2.2 Weather-Aware Pest Management and Urban Pest Data", 2)
    add_para(
        doc,
        "Hiscox et al. (2025) review the impact of adverse weather events on cockroaches and flies and report that increased temperature, extreme rainfall, and flooding can be associated with changes in pest population behaviour and human-pest interaction. This supports the need for weather-aware risk indicators in pest control planning. However, the same literature also shows that causal claims are difficult because pest behaviour depends on many confounding factors such as sanitation, habitat disruption, infrastructure condition, and pesticide use.",
    )
    add_para(
        doc,
        "Richardson et al. (2025) analyse urban rat reports across multiple cities and find that rat numbers are linked with climate warming, urbanization, and population density. While the study is not specific to Malaysian pest control service records, it demonstrates the importance of combining pest reports with environmental and urban context. Brimblecombe et al. (2023) also show that urban pest enquiry databases can reveal long-term patterns, but such records are heterogeneous and may be affected by reporting bias. These studies justify using Nomobug's internal service outcomes carefully as business evidence rather than as universal ecological proof.",
    )

    add_heading(doc, "2.3 Digital Pest Decision Support Systems", 2)
    add_para(
        doc,
        "Pinto et al. (2025) describe the Network for Environment and Weather Applications (NEWA), a weather-based pest and disease decision support platform that integrates historical and forecast weather data with risk models. Although NEWA focuses on agricultural integrated pest management, it demonstrates the value of combining weather data, dashboards, alerts, and decision-support outputs. The proposed Nomobug system adapts this idea to an urban pest-control SME context by combining service outcomes, warranty records, recurrence windows, and local weather indicators.",
    )
    add_para(
        doc,
        "The research gap is that most pest decision-support examples focus on agriculture or public-sector surveillance, while SME pest control operations need analytics around customer service outcomes, warranty costs, technician review indicators, and upsell planning. This project therefore contributes by applying weather-aware decision support to a real commercial service dataset rather than only biological pest observations.",
    )

    add_heading(doc, "2.4 Spatial-Temporal Hotspot Detection", 2)
    add_para(
        doc,
        "Spatial-temporal analytics are useful when the location and timing of events matter. Li et al. (2023) propose a spatiotemporal clustering algorithm for satellite hotspot data and discuss the importance of density-based clustering, arbitrary cluster shapes, noise handling, and parameter selection. Their work supports the use of DBSCAN-style clustering for Nomobug's repeat-problem data because warranty cases may form irregular clusters and random isolated claims should not automatically become hotspots.",
    )
    add_para(
        doc,
        "KDE heatmaps complement clustering because they provide a smoother visual layer for dashboard users. In this project, DBSCAN will be used as the main clustering method for detecting candidate repeat-problem hotspots, while KDE-style heatmaps will be used mainly for visual interpretation. This keeps the method understandable for management while still providing a stronger spatial analysis contribution than a simple area-count table.",
    )

    add_heading(doc, "2.5 Weather-Based Prediction and Explainable Modelling", 2)
    add_para(
        doc,
        "Sumido et al. (2023) demonstrate that pest occurrence can be modelled using weather variables and machine learning techniques, showing that weather can be a meaningful input for pest prediction. However, Nomobug's warranty dataset is expected to contain approximately 200+ warranty records, which is small for complex machine learning. Therefore, this project will use lightweight and explainable methods first, such as time-lag rainfall analysis, simple comparative recurrence rates, and logistic regression only if the feature set and target variable are sufficiently stable.",
    )
    add_para(
        doc,
        "This choice is methodologically important. Complex models such as random forests, XGBoost, or deep learning may appear advanced, but they can overfit small datasets and reduce explainability. A capstone system used by a real company should prioritise trustworthy, understandable indicators over black-box model accuracy that cannot be defended with limited data.",
    )

    add_heading(doc, "2.6 AI-Assisted Recommendations for SMEs", 2)
    add_para(
        doc,
        "Lu et al. (2022) review AI-enabled opportunities and transformation challenges for SMEs, showing that AI can support SME resilience and performance but must be adopted with attention to resource constraints, organisational readiness, and practical value. This directly supports the proposed future use of local Ollama-based recommendations. Rather than making AI the source of truth, the proposed system will calculate analytics-backed scores first and then use local AI to generate plain-language suggestions for management.",
    )
    add_para(
        doc,
        "This approach keeps the AI layer low-cost and privacy-conscious. The company can run Ollama locally in the future to explain why a case is high risk, why an area needs closer monitoring, or which action may reduce repeat visits. This is aligned with prescriptive analytics while avoiding external AI API cost.",
    )

    add_heading(doc, "2.7 Literature Matrix", 2)
    add_table(
        doc,
        ["Study", "Focus", "How It Supports This Project", "Gap Addressed"],
        [
            ("Ragazou et al. (2023)", "BI for SMEs", "Justifies BI as decision support for resource-limited SMEs", "Applies BI to a specific pest-control SME"),
            ("Gurcan et al. (2023)", "BI trends using ML topic modelling", "Supports data mining, visualisation, big data analytics, and AI as current BI directions", "Translates broad BI trends into operational modules"),
            ("Hiscox et al. (2025)", "Weather events and hygiene-related pests", "Supports rainfall/flood/temperature as associated risk indicators", "Avoids causal claims and applies association to service outcomes"),
            ("Richardson et al. (2025)", "Urban rats, warming, urbanization", "Supports combining pest reports with environmental and urban context", "Uses company service outcomes instead of public reports"),
            ("Brimblecombe et al. (2023)", "Urban pest enquiries database", "Shows pest records can reveal temporal patterns but may contain reporting bias", "Adds operational outcomes such as claims/refunds"),
            ("Pinto et al. (2025)", "Weather-based pest DSS", "Supports integrating weather data, alerts, and decision tools", "Adapts DSS idea from agriculture to SME pest services"),
            ("Li et al. (2023)", "Spatiotemporal hotspot clustering", "Supports DBSCAN-style clustering, noise handling, and parameter tuning", "Applies clustering to warranty/repeat-service hotspots"),
            ("Sumido et al. (2023)", "Weather-based pest prediction", "Supports weather variables as model inputs", "Uses lightweight modelling due to small warranty dataset"),
            ("Lu et al. (2022)", "AI adoption in SMEs", "Supports future local AI recommendation layer", "Keeps AI low-cost, explainable, and company-ready"),
        ],
        widths=[1.3, 1.3, 2.0, 1.7],
    )

    add_heading(doc, "2.8 Research Gap", 2)
    add_para(
        doc,
        "The reviewed literature shows that BI can improve SME decision-making, weather can be an important pest-risk indicator, spatial-temporal clustering can detect emerging hotspots, and AI can support SME decision processes. However, there is limited work that combines these areas into a single practical system for an urban pest-control SME using real operational records such as warranty claims, refunds, payment links, Google Calendar service slots, technician assignments, and upsell windows.",
    )
    add_para(
        doc,
        "This project fills that gap by building a company-specific decision-support dashboard that integrates operational, environmental, spatial, and financial data. The contribution is not a heavy machine learning model; it is a usable analytics architecture that helps management understand where problems occur, when recurrence happens, which factors require review, and what actions may improve planning.",
    )


def add_chapter3(doc):
    add_heading(doc, "Chapter 3: Methodology", 1)
    add_heading(doc, "3.1 Overall Research and Development Approach", 2)
    add_para(
        doc,
        "The project will follow an applied data analytics development methodology. The process begins with data understanding and cleaning, then moves into data integration, feature engineering, analytics modelling, dashboard development, and evaluation with company users. The methodology is designed to suit messy SME data and a limited warranty dataset, so explainability and actionability are prioritised over complex black-box machine learning.",
    )
    add_table(
        doc,
        ["Phase", "Main Activities", "Main Output"],
        [
            ("1. Data understanding", "Profile Google Sheets, Calendar data, weather sources, missing values, duplicates, inconsistent labels", "Data quality report and field mapping"),
            ("2. Data preparation", "Clean dates, IDs, phone numbers, areas, pest types, packages, refund categories, calendar event types", "Cleaned staging tables and validation flags"),
            ("3. Data integration", "Join data by customer ID, phone number, service date, area, and calendar event matching", "Analytics-ready fact/dimension model"),
            ("4. Weather enrichment", "Map area/postcode to coordinates and retrieve rainfall, humidity, and temperature windows", "Weather-enriched service table and coverage report"),
            ("5. Analytics modelling", "Compute KPIs, recurrence windows, treatment difficulty, hotspot clusters, upsell package fit, warranty review factors", "Dashboard marts and scoring tables"),
            ("6. Scheduled refresh", "Run Python ETL or BigQuery scheduled queries on a planned interval", "Near-real-time refreshed analytics tables"),
            ("7. Dashboard development", "Build Superset dashboards connected to BigQuery with filters and charts", "Browser-based BI dashboard"),
            ("8. Recommendation logging", "Store generated recommendations, reasons, management response, and later outcome", "Recommendation log for review and future AI learning"),
            ("9. Evaluation", "Review dashboard with management, compare manual vs dashboard analysis, collect feedback", "Evaluation findings and improvement list"),
        ],
        widths=[1.25, 3.2, 1.8],
    )

    add_heading(doc, "3.2 Proposed Technical Architecture", 2)
    add_para(
        doc,
        "The selected architecture is Google Sheets or CSV/Excel as the raw data source, Python as the ETL and analytics engine, BigQuery as the cloud analytics warehouse, Apache Superset as the visualization layer, and optional local Ollama as a future AI recommendation layer. Docker will be used for local reproducibility where suitable, especially for Superset, the Python worker, and future Ollama testing.",
    )
    add_para(
        doc,
        "The workflow is designed as a controlled analytics pipeline. Raw data is first extracted from Google Sheets, CSV/Excel exports, and Google Calendar. The Python ETL process validates the records, standardises field values, flags data quality issues, enriches services with weather and location features, and loads the output into BigQuery. BigQuery then stores three layers: raw staging tables, cleaned core tables, and dashboard-ready marts. Superset reads only the cleaned core or mart tables so users do not need to interact with raw inconsistent records.",
    )
    add_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ("Data source", "Google Sheets, CSV/Excel, Google Calendar", "Company-facing raw data entry and operational records"),
            ("ETL and analytics", "Python, Pandas, scikit-learn, BigQuery client", "Cleaning, enrichment, feature engineering, clustering, scoring"),
            ("Cloud warehouse", "Google BigQuery", "Store raw, cleaned, and dashboard-ready analytics tables"),
            ("Visualization", "Apache Superset", "Browser-based dashboards, charts, filters, maps, and tables"),
            ("Refresh pipeline", "Scheduled Python ETL or BigQuery scheduled queries", "Daily or weekly near-real-time refresh of analytics tables"),
            ("Governance", "Anonymised views and validation logs", "Protect customer data and show data quality status"),
            ("AI recommendation", "Local Ollama (future phase)", "Generate plain-language suggestions from analytics outputs without paid API calls"),
            ("Deployment", "Docker Compose for local services; BigQuery hosted by Google Cloud", "Reproducible setup for demo and company handover"),
        ],
        widths=[1.35, 2.0, 2.95],
    )
    add_callout(
        doc,
        "Architecture flow",
        "Google Sheets / CSV / Calendar -> Python ETL -> BigQuery staging, core, and mart tables -> Superset dashboards -> management decisions. Optional future flow: BigQuery scores -> local AI service -> recommendation log -> dashboard.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "3.3 BigQuery Data Warehouse Design", 2)
    add_para(
        doc,
        "BigQuery is selected because the project is analytics-oriented and the company data is already connected to the Google ecosystem. BigQuery removes the need to maintain a database server while still allowing SQL-based modelling, scheduled queries, partitioned tables, and dashboard connections. The warehouse will be organised into three logical layers.",
    )
    add_table(
        doc,
        ["Layer", "Purpose", "Example Tables"],
        [
            ("Staging", "Store raw imported data with minimal transformation", "stg_customers, stg_services, stg_warranty_claims, stg_refunds, stg_calendar_events"),
            ("Cleaned core", "Standardised entities and facts with consistent keys", "dim_customer, dim_area, dim_pest, fact_service, fact_payment, fact_warranty, fact_refund"),
            ("Analytics marts", "Dashboard-ready tables and views", "mart_area_risk, mart_treatment_difficulty, mart_recurrence_windows, mart_upsell_fit, mart_marketing_funnel"),
        ],
        widths=[1.2, 2.7, 2.4],
    )
    add_para(
        doc,
        "The design will use stable identifiers such as customer_id, service_id, date, area, postcode, pest_type, package_type, technician, and team. BigQuery is used as an analytics warehouse rather than a transaction-processing database, so relationships will be managed through consistent keys, SQL joins, views, and mart-building queries rather than application-level transactional foreign-key enforcement. Where direct keys are missing, matching will use a controlled fallback sequence: customer ID, normalised phone number, and then customer name plus service/calendar date as a lower-confidence match.",
    )
    add_table(
        doc,
        ["Relationship", "Join Key / Matching Rule", "Purpose"],
        [
            ("dim_customer -> fact_service", "customer_id; fallback to normalised phone number", "Connect services to customer profile, premise type, and segment"),
            ("fact_service -> fact_warranty", "service_id when available; otherwise customer_id + service date window + pest type", "Link claims to original service context"),
            ("fact_service -> fact_refund", "customer_id/service_id and refund date after service date", "Measure refund impact by service, area, pest, and technician review signal"),
            ("fact_service -> fact_calendar_event", "service date + customer name/phone/event title; confidence flag stored", "Compare scheduled slots with service and warranty activity"),
            ("fact_service -> dim_area", "standardised area/postcode", "Attach city/state, centroid coordinates, and area match confidence"),
            ("fact_service -> weather features", "area/postcode coordinate + service date", "Attach rainfall, humidity, temperature, and weather data coverage"),
            ("fact_upsell -> fact_payment", "customer_id + payment date near offer/window date", "Infer window performance when direct response data is unavailable"),
            ("mart_recommendations -> facts/marts", "related customer_id, service_id, area, pest type, and generated date", "Track suggested actions and later outcomes"),
        ],
        widths=[2.0, 2.45, 1.85],
    )

    add_heading(doc, "3.4 Data Sources and Field Mapping", 2)
    add_table(
        doc,
        ["Data Source", "Key Fields", "Analytics Use"],
        [
            ("Prospects / leads", "phone, lead date, ad source, campaign, status, PIC, follow-up count", "Marketing attribution, lead conversion, funnel analysis"),
            ("Services / sales", "customer ID, service date, pest type, package, area, amount, payment status", "Service trends, revenue, customer segmentation, recurrence base"),
            ("Warranty claims", "claim date, complaint reason, technician/team, repeat flag, complimentary service", "Warranty review factors, repeat-problem analysis, treatment difficulty"),
            ("Refunds", "refund date, reason, amount, status", "Root cause categories, financial loss, difficulty score"),
            ("Payments and payment links", "received date, amount, method, sent date, won/pending", "Payment completion and payment-link conversion"),
            ("Upsell records", "upsell package, offer date/window, accepted/purchased package, amount", "Package effectiveness and package-fit matrix"),
            ("Google Calendar", "event title, start/end time, area, technician/team, session type", "Capacity analytics, slot utilisation, warranty/extra session share"),
            ("Weather/flood data", "rainfall, humidity, temperature, warning/flood flags, date, area coordinate", "Weather-associated risk, service-date enrichment"),
        ],
        widths=[1.35, 2.4, 2.55],
    )

    add_heading(doc, "3.5 Data Cleaning, Inclusion, and Exclusion Rules", 2)
    add_para(
        doc,
        "The ETL process will not simply load the company files directly into dashboards. Each source will pass through validation and standardisation rules before it is used for analysis. Raw data will be preserved in staging tables, while cleaned values and quality flags will be stored in the cleaned core layer. This allows the project to explain which records were included, corrected, flagged for review, or excluded from a specific analysis.",
    )
    add_table(
        doc,
        ["Data Element", "Cleaning / Standardisation Rule", "Flag or Exclusion Rule"],
        [
            ("Dates", "Convert all service, payment, claim, refund, upsell, and calendar dates into a consistent date format", "Invalid or impossible dates are flagged and excluded from time-based analysis until corrected"),
            ("Customer identifiers", "Standardise customer_id and normalise phone numbers by removing spaces, symbols, and country-code inconsistencies", "Rows without any usable customer_id, phone, or name/date fallback are flagged as unmatched"),
            ("Pest type", "Map variants such as lipas, cockroach, cockroaches, anai-anai, and termite into standard pest labels", "Unknown pest labels are retained but flagged for manual mapping"),
            ("Area/postcode", "Standardise area names and map postcode/area to centroid coordinates", "Unknown areas are included in non-map summaries but excluded from coordinate-based hotspot analysis"),
            ("Technician/team", "Standardise spelling and team labels across service, warranty, and calendar records", "Missing technician is allowed for company-level analysis but flagged for technician review dashboards"),
            ("Package/service type", "Standardise 1x, 2x, 3x, yearly, TRBS, GPC, termite, and related package labels", "Unrecognised package labels are flagged and grouped as Other until reviewed"),
            ("Refund reason", "Group free-text reasons into categories such as Cannot eradicate, Not satisfied, Billing issue, Late service, and Manual review", "Ambiguous reasons are shown in root-cause analysis as Needs manual review"),
            ("Calendar event", "Parse event title/description into New/Paid, Warranty, Extra, Follow-up, or Unknown", "Unmatched or unknown events are kept in scheduling totals but excluded from service-quality attribution"),
            ("Weather match", "Attach service-date rainfall, humidity, and temperature using area/postcode centroid", "Rows without weather match are kept with weather_data_status = Missing and excluded from weather-specific charts"),
        ],
        widths=[1.35, 2.75, 2.2],
    )
    add_para(
        doc,
        "The project will use a conservative exclusion policy. Records are not silently deleted. They are first preserved in staging tables, then marked with validation flags such as invalid_date, duplicate_candidate, unmatched_customer, unknown_area, missing_weather, or low_calendar_match_confidence. A record is excluded only from the analysis that cannot safely use it. For example, a service with missing coordinates can still be used in revenue trends, but not in DBSCAN hotspot detection.",
    )
    add_table(
        doc,
        ["Data Quality Metric", "Purpose"],
        [
            ("Missing-value summary", "Shows which fields are incomplete by source and by column"),
            ("Duplicate-record count", "Identifies repeated customers, services, payments, claims, or calendar rows"),
            ("Invalid-date count", "Prevents broken time trends and recurrence calculations"),
            ("Unmatched-customer count", "Shows records that cannot be joined confidently across sheets"),
            ("Area/postcode match confidence", "Indicates whether spatial analysis is based on strong or approximate location matching"),
            ("Weather-enrichment coverage", "Shows the percentage of service records with matched rainfall/humidity/temperature"),
            ("Calendar matching confidence", "Shows how reliable the service-to-calendar linkage is"),
        ],
        widths=[2.3, 4.0],
    )

    add_heading(doc, "3.6 Analytics Methods", 2)
    doc.paragraphs[-1].insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
    add_table(
        doc,
        ["Question", "Method", "Output"],
        [
            ("What happened?", "Descriptive BI by day/week/month", "KPIs, trends, ranking tables, distributions"),
            ("Where are repeat problems concentrated?", "DBSCAN-style spatial clustering plus KDE heatmap visualization", "Hotspot clusters and density map"),
            ("When do pest issues recur?", "Days-to-recurrence and interval comparison", "Average/median recurrence windows and recurrence buckets"),
            ("Is weather associated with higher warranty/repeat risk?", "Rainfall time-lag features, correlation review, and logistic regression if data is sufficient", "Weather-associated indicators and risk probabilities"),
            ("Which areas are harder to treat?", "Weighted treatment difficulty index", "Area/pest/team difficulty level and planning action"),
            ("Which package fits which customer group?", "Frequency and revenue-based package-fit matrix", "Recommended upsell package by customer/pest/area segment"),
            ("What should management do?", "Prescriptive scoring rules and optional local AI explanation", "Recommended follow-up, assignment, review, or upsell action"),
        ],
        widths=[1.85, 2.6, 1.85],
    )

    add_heading(doc, "3.7 Key Analytics Logic", 2)
    add_para(doc, "Treatment Difficulty Index:", bold_lead="Treatment Difficulty Index:")
    add_para(
        doc,
        "The treatment difficulty index will combine repeat-problem rate, warranty-claim rate, refund rate, complimentary-service rate, and extra-session rate. The initial weights will be reviewed with management because the score is intended as a planning indicator, not an absolute measurement.",
    )
    add_para(doc, "Weather windows:", bold_lead="Weather windows:")
    add_para(
        doc,
        "For each service record, the ETL process will calculate rain_same_day_mm, rain_previous_1_day_mm, rain_last_3_days_mm, rain_last_7_days_mm, temperature_mean, humidity_mean, and heavy_rain_flag. Warranty and repeat-problem analysis will use service-date weather windows, while area risk dashboards may also show recent/current weather indicators.",
    )
    add_para(doc, "Warranty factor categories:", bold_lead="Warranty factor categories:")
    add_para(
        doc,
        "Warranty claims will be grouped into review categories: Weather-associated, Technician or service quality review, Area/premise condition, Customer expectation, and Needs manual review. The wording will remain careful and non-punitive.",
    )
    add_para(doc, "Recurrence buckets:", bold_lead="Recurrence buckets:")
    add_para(
        doc,
        "Recurrence windows will be grouped into 0-7 days, 8-14 days, 15-30 days, 31-60 days, and 61+ days. These buckets help management plan follow-ups around observed recurrence patterns.",
    )
    add_para(doc, "Scheduled refresh:", bold_lead="Scheduled refresh:")
    add_para(
        doc,
        "The system will use a near-real-time refresh design rather than streaming real-time processing. A scheduled Python ETL job or BigQuery scheduled query will refresh cleaned tables and dashboard marts daily or weekly, depending on company needs and data update frequency. This is sufficient for management reporting, warranty review, and service planning because the main business questions are based on day/week/month trends rather than second-by-second operations.",
    )
    add_para(doc, "Recommendation log:", bold_lead="Recommendation log:")
    add_para(
        doc,
        "Prescriptive outputs will be stored in a recommendation log. Each record will include recommendation type, generated date, related customer/service/area, supporting reason, confidence or review label, management action taken, and later outcome where available. This allows the company to review whether suggested actions were useful and prepares the system for a future Ollama-based AI explanation layer.",
    )

    add_heading(doc, "3.8 Dashboard Design", 2)
    add_para(
        doc,
        "The visualization layer will be implemented in Apache Superset to reduce custom frontend development and align the project with a data analytics capstone. Superset will connect to BigQuery tables/views and provide dashboards with filters for date range, grain, area, pest type, package type, team, technician, customer segment, and session type.",
    )
    add_table(
        doc,
        ["Dashboard Area", "Main Visuals"],
        [
            ("Executive Summary", "KPI cards, service/revenue/warranty trends, top areas, active risk indicators"),
            ("Marketing and Funnel", "Lead source quality, conversion funnel, follow-up depth, PIC performance"),
            ("Area and Weather Risk", "Map, rainfall indicators, area risk ranking, weather-associated labels"),
            ("Treatment Difficulty", "Difficulty ranking, area/pest/team score table, recommended planning actions"),
            ("Warranty and Repeat Analysis", "Warranty factors by area/pest/technician, review queue, repeat rate trends"),
            ("Recurrence and Hotspots", "Recurrence buckets, DBSCAN clusters, KDE map layer, cluster movement by period"),
            ("Upsell Intelligence", "Package effectiveness, package-fit matrix, window performance"),
            ("Scheduling Capacity", "Slots used, new/warranty/extra split, peak days/times, area density"),
            ("Refund Root Cause", "Refund categories, amount/rate by area, pest, technician, and source"),
            ("Data Quality and Refresh", "Missing values, duplicates, invalid dates, match confidence, weather coverage, refresh status"),
            ("Recommendation Log", "Suggested action, reason, review label, management response, and outcome tracking"),
        ],
        widths=[2.0, 4.3],
    )

    add_heading(doc, "3.9 Data Governance and Privacy", 2)
    add_para(
        doc,
        "The dashboard will separate internal operational data from academic/reporting views. Sensitive fields such as customer names, phone numbers, and full addresses will be available only where needed for internal matching. Management dashboards and academic outputs will use customer IDs, general area/postcode, pest type, service type, and aggregated metrics. Technician-related analytics will be framed as review indicators for quality improvement rather than blame or direct causation.",
    )
    add_para(
        doc,
        "The BigQuery model will support anonymised views for reporting. These views will remove or mask personally identifiable information while preserving the fields needed for analysis, such as service date, area, pest type, package type, claim status, refund category, and difficulty level. This supports privacy, ethics, and company confidence in using the system.",
    )

    add_heading(doc, "3.10 AI Recommendation Extension", 2)
    add_para(
        doc,
        "A future enhancement will connect a local Ollama model to the analytics system. The recommended design is to keep the AI outside BigQuery and Superset as a local service. The Python backend or worker will retrieve relevant scores from BigQuery, pass structured context to Ollama, and store the generated recommendation text back into the recommendation log for dashboard display.",
    )
    add_para(
        doc,
        "This design avoids external AI API cost and protects company data. The AI will explain analytics-backed recommendations such as assigning an experienced technician, scheduling a follow-up within a specific window, reviewing an area/premise condition, or recommending a package-fit upsell. It will not make final decisions or invent unsupported claims.",
    )

    add_heading(doc, "3.11 Computing Factors and Constraints", 2)
    add_table(
        doc,
        ["Factor", "Consideration", "Proposed Handling"],
        [
            ("Performance", "Dashboard queries should remain responsive", "Use BigQuery marts and scheduled transformations rather than recalculating every chart live"),
            ("Security and privacy", "Customer names, phone numbers, and locations are sensitive", "Use access control, anonymised reports, customer IDs, and area/postcode-level coordinates"),
            ("Ethics", "Technician and weather analysis can be misread as blame or causation", "Use wording such as review indicator, associated factor, and requires manual review"),
            ("Usability", "Management needs quick insights, not technical complexity", "Use KPI cards, maps, filters, ranking tables, and short recommendations"),
            ("Data quality", "Sheets may contain duplicates, missing fields, inconsistent labels", "Create cleaning rules, validation reports, unmatched-record logs, and data quality dashboard"),
            ("Refresh reliability", "Company data changes over time", "Use scheduled ETL/queries and show last refresh time and failed refresh status"),
            ("Cost", "Company prefers low-cost tools", "Use BigQuery within free/low usage limits, Superset open source, Open-Meteo, and local Ollama later"),
            ("Model reliability", "Warranty dataset is limited", "Prefer explainable scoring and simple statistical methods before complex ML"),
        ],
        widths=[1.4, 2.4, 2.5],
    )

    add_heading(doc, "3.12 Evaluation Plan", 2)
    for item in [
        "Functional testing: verify data upload/sync, cleaning outputs, BigQuery table creation, and dashboard refresh.",
        "Data quality testing: verify missing-value summaries, duplicate counts, invalid-date flags, unmatched-customer logs, coordinate match confidence, weather coverage, and calendar matching confidence.",
        "Refresh testing: confirm scheduled ETL or BigQuery scheduled queries update marts and show last refresh status.",
        "Analytics validation: compare selected dashboard metrics against manual spreadsheet calculations.",
        "Usability testing: ask management or operations users to complete business questions using the dashboard.",
        "Privacy testing: confirm anonymised views remove customer names, phone numbers, and full address fields from academic/reporting outputs.",
        "Recommendation testing: verify recommendation log fields are generated and that management response/outcome can be tracked.",
        "Insight usefulness: collect feedback on whether area risk, recurrence, treatment difficulty, and upsell outputs are actionable.",
        "Method review: confirm that weather and technician analysis is framed as association/review indicator, not direct blame or causation.",
    ]:
        add_bullet(doc, item)


def add_chapter4(doc):
    add_heading(doc, "Chapter 4: Work Plan and Timeline", 1)
    add_para(
        doc,
        "The project will be completed in iterative stages. The timeline below assumes a 14-week capstone implementation period after proposal approval. Activities are arranged to reduce risk: data understanding and warehouse design come first, then analytics logic, then dashboards, then evaluation.",
    )
    add_table(
        doc,
        ["Week", "Work Package", "Deliverable"],
        [
            ("1", "Confirm scope, finalise data sources, review rubric and supervisor feedback", "Approved project scope and data access checklist"),
            ("2", "Data profiling for Google Sheets, warranty records, refunds, services, payments, and calendar exports", "Data quality report"),
            ("3", "Design BigQuery staging/core/mart schema, table relationships, matching keys, anonymised views, and recommendation log", "Warehouse design document"),
            ("4", "Build Python ETL for data import, standardisation, validation flags, and inclusion/exclusion reporting", "Cleaned staging pipeline and data quality dashboard source"),
            ("5", "Implement area/postcode lookup, weather enrichment features, and weather coverage reporting", "Weather-enriched service table"),
            ("6", "Build descriptive analytics marts for dashboard overview, service, sales, warranty, refund, and scheduling", "Core BI marts"),
            ("7", "Develop treatment difficulty, recurrence window, and warranty factor analytics", "Quality and recurrence marts"),
            ("8", "Develop DBSCAN hotspot detection and KDE map-ready outputs", "Hotspot output table"),
            ("9", "Develop upsell package effectiveness, package-fit, window performance analytics, and recommendation log outputs", "Upsell intelligence and recommendation marts"),
            ("10", "Connect Apache Superset to BigQuery and build first dashboard pages", "Dashboard prototype v1"),
            ("11", "Refine filters, maps, data quality views, refresh status, rankings, and recommendations with company feedback", "Dashboard prototype v2"),
            ("12", "Evaluate dashboard against manual spreadsheet answers and collect usability feedback", "Evaluation report draft"),
            ("13", "Prepare final report, screenshots, architecture diagram, data flow diagram, and viva explanation", "Final documentation draft"),
            ("14", "Final testing, polishing, supervisor review, and submission packaging", "Final project submission"),
        ],
        widths=[0.65, 3.75, 1.9],
    )

    add_heading(doc, "4.1 Milestones", 2)
    add_table(
        doc,
        ["Milestone", "Target Week", "Success Criteria"],
        [
            ("M1: Data source confirmation", "Week 1", "Required sheets/calendar exports identified and accessible"),
            ("M2: Clean analytics dataset", "Week 4", "Cleaning rules run successfully and data quality issues are documented"),
            ("M3: BigQuery warehouse ready", "Week 6", "Staging, core, and mart tables created and populated"),
            ("M4: Core analytics complete", "Week 9", "Treatment, recurrence, hotspot, weather, and upsell analytics produce outputs"),
            ("M5: Dashboard prototype complete", "Week 11", "Superset dashboards answer key management questions including data quality and refresh status"),
            ("M6: Evaluation complete", "Week 12", "Management feedback and manual-check comparison documented"),
            ("M7: Final submission ready", "Week 14", "Report, dashboard screenshots, and demonstration environment prepared"),
        ],
        widths=[2.0, 1.0, 3.3],
    )

    add_heading(doc, "4.2 Risks and Mitigation", 2)
    doc.paragraphs[-1].insert_paragraph_before().add_run().add_break(WD_BREAK.PAGE)
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ("Warranty data has only around 200+ records", "Complex ML may overfit", "Use explainable scoring, lag analysis, and logistic regression only if stable"),
            ("Google Sheets columns change", "ETL can break", "Use canonical column mapping and alias handling"),
            ("Calendar event names are inconsistent", "Session classification errors", "Use keyword rules and manual review list"),
            ("Coordinates are incomplete", "Hotspot analysis becomes less precise", "Use postcode/area centroid fallback and document confidence"),
            ("Weather data does not match micro-location conditions", "Risk of overclaiming weather effect", "Use association wording and service-date weather windows"),
            ("Scheduled refresh fails or Google Sheets access changes", "Dashboard may show stale data", "Show last successful refresh, failure status, and manual refresh fallback"),
            ("BigQuery/Superset setup takes longer than expected", "Dashboard delivery risk", "Build static CSV fallback dashboards if connection issues occur"),
            ("Company users need simpler explanations", "Low adoption", "Use short recommendation labels and management-friendly dashboards"),
        ],
        widths=[2.0, 1.65, 2.65],
    )


def add_appendices(doc):
    add_heading(doc, "Appendix A: Proposed Dashboard Modules", 1)
    for module in [
        ("Executive Summary", "Overview of customers, services, revenue, warranty claims, refunds, recurrence, weather-associated risk, hotspots, and upsell performance."),
        ("Marketing and Funnel", "Ad source quality, prospect conversion, follow-up depth, PIC performance, payment-link conversion, and lead-to-close time."),
        ("Area and Weather Risk", "Area-level risk indicators using rainfall windows, flood/warning flags, pest volume, repeat rate, and complaint/refund signals."),
        ("Treatment Difficulty", "Ranking of area, pest type, and team based on repeat problems, warranty claims, refunds, complimentary services, and extra sessions."),
        ("Warranty and Repeat Analysis", "Warranty factor categories by period, area, pest type, team, and technician with a review queue."),
        ("Recurrence and Hotspots", "Days-to-recurrence buckets and DBSCAN/KDE hotspot map for emerging repeat-problem clusters."),
        ("Upsell Intelligence", "Package effectiveness, package-fit matrix, upsell window performance, and future response tracking template."),
        ("Scheduling Capacity", "Calendar slot usage, new/warranty/extra split, peak booking times, area density, and session gap analysis."),
        ("Refund Root Cause", "Standardised refund categories and financial impact by area, pest, source, and technician review indicator."),
        ("Data Quality and Refresh", "Missing-value summary, duplicate count, invalid-date records, unmatched customer records, coordinate match confidence, weather enrichment coverage, calendar match confidence, last refresh time, and failed refresh status."),
        ("Recommendation Log", "Generated prescriptive suggestions, supporting reasons, confidence or review labels, management action taken, and later outcome tracking for continuous improvement."),
    ]:
        add_heading(doc, module[0], 3)
        add_para(doc, module[1])

    add_heading(doc, "Appendix B: Proposed Data Warehouse Tables", 1)
    add_table(
        doc,
        ["Table/View", "Type", "Purpose"],
        [
            ("stg_* tables", "Raw staging", "Preserve imported Google Sheets/CSV/calendar data with source metadata"),
            ("dim_customer", "Dimension", "Customer profile with anonymised reporting support"),
            ("dim_area", "Dimension", "Area/postcode/city/state and centroid coordinates"),
            ("dim_pest", "Dimension", "Standardised pest type labels"),
            ("fact_service", "Fact", "One row per service/session with date, area, pest, package, technician/team"),
            ("fact_warranty", "Fact", "Warranty claims, suspected factors, review priority, repeat/unresolved flags"),
            ("fact_refund", "Fact", "Refund amounts, reasons, categories, and statuses"),
            ("fact_payment", "Fact", "Payment received and payment-link conversion records"),
            ("fact_upsell", "Fact", "Upsell package, offer/window date, accepted/purchased indicators"),
            ("fact_calendar_event", "Fact", "Scheduled slots, duration, parsed session type, area/team match"),
            ("fact_recommendation_log", "Fact", "Recommendation type, generated date, related customer/service/area, reason, review label, action taken, and outcome"),
            ("dq_validation_log", "Quality table", "Missing values, duplicate candidates, invalid dates, unmatched records, and source-level quality status"),
            ("dq_match_confidence", "Quality table", "Customer, calendar, area/postcode, and weather-enrichment match confidence"),
            ("privacy_reporting_views", "Views", "Anonymised reporting outputs using customer ID and area/postcode instead of names/phone numbers"),
            ("mart_area_risk", "Analytics mart", "Area weather-associated risk score and supporting factors"),
            ("mart_treatment_difficulty", "Analytics mart", "Difficulty index by area, pest type, and team"),
            ("mart_recurrence_windows", "Analytics mart", "Days-to-recurrence and bucketed recurrence outcomes"),
            ("mart_hotspots", "Analytics mart", "DBSCAN cluster IDs, centroid, severity, and cluster trend"),
            ("mart_upsell_fit", "Analytics mart", "Best-fit package recommendations by customer/pest/area segment"),
            ("mart_recommendations", "Analytics mart", "Dashboard-ready prescriptive recommendations and optional AI-generated explanation text"),
            ("mart_refresh_status", "Analytics mart", "Last successful refresh, failed refresh reason, row counts, and dashboard data freshness"),
        ],
        widths=[2.1, 1.3, 2.9],
    )


def add_references(doc):
    add_heading(doc, "References", 1)
    refs = [
        "Apache Superset. (2026). Google BigQuery database support. https://superset.apache.org/user-docs/databases/supported/google-bigquery",
        "Apache Superset. (2026). Introduction to Apache Superset. https://superset.apache.org/user-docs/intro/",
        "Brimblecombe, P., Müller, G., Schmidt, M., Tischhauser, W., Landau, I., & Querner, P. (2023). Urban pest abundance and public enquiries in Zurich 1991-2022. Insects, 14(10), 798. https://doi.org/10.3390/insects14100798",
        "Google Cloud. (2026). BigQuery overview. https://cloud.google.com/bigquery",
        "Google Cloud. (2026). BigQuery documentation. https://cloud.google.com/bigquery/docs",
        "Gurcan, F., Ayaz, A., Menekse Dalveren, G. G., & Derawi, M. (2023). Business intelligence strategies, best practices, and latest trends: Analysis of scientometric data from 2003 to 2023 using machine learning. Sustainability, 15(13), 9854. https://doi.org/10.3390/su15139854",
        "Hiscox, A., Spencer, F., Dennehy, J., Dyall, W., Jenkins, A., Narendran, A., Das, A., Logan, J. G., & Jones, R. T. (2025). The impact of adverse weather events on cockroaches and flies, and the possible effects on disease. Medical and Veterinary Entomology, 39(3), 500-514. https://doi.org/10.1111/mve.12797",
        "Li, W., Dodwell, E., & Cook, D. (2023). A clustering algorithm to organize satellite hotspot data for the purpose of tracking bushfires remotely. The R Journal, 15(1), 17-33. https://journal.r-project.org/articles/RJ-2023-022/",
        "Lu, X., Wijayaratna, K., Huang, Y., & Qiu, A. (2022). AI-enabled opportunities and transformation challenges for SMEs in the post-pandemic era: A review and research agenda. Frontiers in Public Health, 10, 885067. https://doi.org/10.3389/fpubh.2022.885067",
        "Open-Meteo. (2026). Historical Weather API. https://open-meteo.com/en/docs/historical-weather-api",
        "Pinto, A. F., Olmstead, D., Calixto, A. A., & Gómez, M. I. (2025). Network for Environment and Weather Applications: An overview of the digital pest management decision support tool. Applied Economics Teaching Resources. https://www.aetrjournal.org/UserFiles/file/AETR_2025_0218%20Final.pdf",
        "Ragazou, K., Passas, I., Garefalakis, A., & Zopounidis, C. (2023). Business intelligence model empowering SMEs to make better decisions and enhance their competitive advantage. Discover Analytics, 1, 2. https://doi.org/10.1007/s44257-022-00002-3",
        "Richardson, J. L., McCoy, E. P., Parlavecchio, N., et al. (2025). Increasing rat numbers in cities are linked to climate warming, urbanization, and human population. Science Advances, 11(5), eads6782. https://doi.org/10.1126/sciadv.ads6782",
        "Sumido, E. C., Feliscuzo, L. S., & Aliac, C. J. G. (2023). Pest classification and prediction: Analyzing the impact of weather to pest occurrence through machine learning. Journal of Engineering Science and Technology, 18, 124-138.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(6)
        p.add_run(ref)


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Nomobug Capstone Proposal")
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED


def main():
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    add_abstract(doc)
    add_toc(doc)
    add_chapter1(doc)
    add_chapter2(doc)
    add_chapter3(doc)
    add_chapter4(doc)
    add_references(doc)
    add_appendices(doc)
    add_footer(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
